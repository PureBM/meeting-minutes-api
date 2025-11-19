import base64
import io
import os
import csv
from datetime import date
from docx import Document

def handler(request):
    try:
        # Parse JSON request body
        body = request.get_json()
        minutes_text = body.get("minutes", "")
        actions = body.get("actions", [])
        title = body.get("title", "Meeting")

        # Template path inside the deployed Vercel environment
        template_path = os.path.join(os.path.dirname(__file__), "Minutes_Template.dotx")

        if os.path.exists(template_path):
            doc = Document(template_path)
        else:
            doc = Document()
            doc.add_paragraph("(Template missing – using default layout)")

        doc.add_heading(f"{title} – {date.today()}", 0)
        doc.add_paragraph(minutes_text)

        # Word file
        word_stream = io.BytesIO()
        doc.save(word_stream)
        word_b64 = base64.b64encode(word_stream.getvalue()).decode("utf-8")

        # CSV file
        csv_stream = io.StringIO()
        writer = csv.writer(csv_stream)
        writer.writerow(["Action Agreed in Detail", "Owner", "Due Date"])
        for a in actions:
            writer.writerow([
                a.get("detail", ""), a.get("owner", ""), a.get("due", "")
            ])
        csv_b64 = base64.b64encode(csv_stream.getvalue().encode("utf-8")).decode("utf-8")

        # ✅ Return structured file objects for GPT to embed
        return {
            "statusCode": 200,
            "headers": {"Content-Type": "application/json"},
            "body": {
                "minutes_file": {
                    "filename": f"Minutes – {title} – {date.today()}.docx",
                    "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "data_base64": word_b64
                },
                "actions_file": {
                    "filename": f"Actions – {title} – {date.today()}.csv",
                    "mime_type": "text/csv",
                    "data_base64": csv_b64
                }
            }
        }

    except Exception as e:
        return {
            "statusCode": 500,
            "body": {"error": str(e)}
        }

